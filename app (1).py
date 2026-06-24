import streamlit as st
import google.generativeai as genai
import time
import io
import zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
#  PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="مولّد الاختبارات الذكي | الجزائر",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  CUSTOM CSS  (iPhone-style, smooth, animated)
# ─────────────────────────────────────────────
st.markdown("""
<style>
/* ── Google Fonts ── */
@import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800&display=swap');

/* ── Root variables ── */
:root {
    --primary:   #4F6EF7;
    --secondary: #7C3AED;
    --accent:    #F59E0B;
    --success:   #10B981;
    --bg:        #F0F4FF;
    --card-bg:   #FFFFFF;
    --text:      #1E293B;
    --sub:       #64748B;
    --radius:    18px;
    --shadow:    0 4px 24px rgba(79,110,247,.10);
    --shadow-hover: 0 12px 36px rgba(79,110,247,.22);
    --transition: all .3s cubic-bezier(.4,0,.2,1);
}

/* ── Global ── */
html, body, [class*="css"] {
    font-family: 'Tajawal', system-ui, -apple-system, 'SF Pro Display',
                 'Segoe UI', sans-serif !important;
    direction: rtl;
    background: var(--bg) !important;
    color: var(--text);
}

/* ── Streamlit chrome ── */
#MainMenu, footer, header {visibility: hidden;}
.block-container { padding: 1.5rem 2rem; max-width: 1100px; }

/* ── Hero banner ── */
.hero {
    background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
    border-radius: var(--radius);
    padding: 2.5rem 2rem;
    text-align: center;
    color: #fff;
    margin-bottom: 2rem;
    box-shadow: var(--shadow-hover);
    animation: fadeSlideDown .6s ease both;
}
.hero h1  { font-size: 2.1rem; font-weight: 800; margin: 0 0 .4rem; }
.hero p   { font-size: 1.05rem; opacity: .88; margin: 0; }

/* ── Cards ── */
.card {
    background: var(--card-bg);
    border-radius: var(--radius);
    padding: 1.6rem 1.8rem;
    box-shadow: var(--shadow);
    margin-bottom: 1.2rem;
    transition: var(--transition);
    border: 1.5px solid transparent;
}
.card:hover {
    transform: translateY(-5px);
    box-shadow: var(--shadow-hover);
    border-color: rgba(79,110,247,.18);
}
.card h3 {
    font-size: 1.1rem;
    font-weight: 700;
    color: var(--primary);
    margin: 0 0 .9rem;
    display: flex;
    align-items: center;
    gap: .45rem;
}

/* ── Streamlit widgets inside cards ── */
.stSelectbox label, .stTextArea label { font-weight: 600; color: var(--text); }
.stSelectbox > div > div,
.stTextArea  > div > div > textarea {
    border-radius: 12px !important;
    border: 1.5px solid #E2E8F0 !important;
    transition: var(--transition);
}
.stSelectbox > div > div:hover,
.stTextArea  > div > div > textarea:focus {
    border-color: var(--primary) !important;
    box-shadow: 0 0 0 3px rgba(79,110,247,.12) !important;
}

/* ── Generate button ── */
div[data-testid="stButton"] > button[kind="primary"] {
    background: linear-gradient(135deg, var(--primary), var(--secondary)) !important;
    color: #fff !important;
    border: none !important;
    border-radius: 14px !important;
    padding: .9rem 2.2rem !important;
    font-size: 1.15rem !important;
    font-weight: 700 !important;
    width: 100% !important;
    box-shadow: 0 6px 24px rgba(79,110,247,.35) !important;
    transition: var(--transition) !important;
    letter-spacing: .5px;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
    transform: translateY(-3px) scale(1.015) !important;
    box-shadow: 0 14px 36px rgba(79,110,247,.45) !important;
}
div[data-testid="stButton"] > button[kind="primary"]:active {
    transform: translateY(0) scale(.98) !important;
}

/* ── Download buttons ── */
div[data-testid="stDownloadButton"] > button {
    border-radius: 12px !important;
    border: 2px solid var(--primary) !important;
    color: var(--primary) !important;
    background: #fff !important;
    font-weight: 600 !important;
    transition: var(--transition) !important;
    width: 100% !important;
}
div[data-testid="stDownloadButton"] > button:hover {
    background: var(--primary) !important;
    color: #fff !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 18px rgba(79,110,247,.25) !important;
}

/* ── Result cards ── */
.result-card {
    background: #fff;
    border-radius: var(--radius);
    padding: 1.8rem;
    margin-bottom: 1.4rem;
    box-shadow: var(--shadow);
    border-right: 5px solid var(--primary);
    animation: fadeSlideUp .5s ease both;
    white-space: pre-wrap;
    font-size: .97rem;
    line-height: 2;
    direction: rtl;
}
.result-card:nth-child(2) { border-right-color: var(--secondary); }
.result-card:nth-child(3) { border-right-color: var(--accent);    }

.result-label {
    display: inline-block;
    background: linear-gradient(135deg, var(--primary), var(--secondary));
    color: #fff;
    padding: .28rem .9rem;
    border-radius: 30px;
    font-size: .85rem;
    font-weight: 700;
    margin-bottom: .8rem;
}

/* ── Progress ── */
.stProgress > div > div > div > div {
    background: linear-gradient(90deg, var(--primary), var(--secondary)) !important;
    border-radius: 8px !important;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1E293B 0%, #0F172A 100%) !important;
    color: #fff;
}
[data-testid="stSidebar"] * { color: #CBD5E1 !important; }
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: #fff !important; font-weight: 700 !important; }

.stat-box {
    background: rgba(255,255,255,.07);
    border: 1px solid rgba(255,255,255,.12);
    border-radius: 14px;
    padding: 1.1rem;
    text-align: center;
    margin: .8rem 0;
}
.stat-number { font-size: 2.4rem; font-weight: 800; color: var(--accent) !important; }
.stat-label  { font-size: .82rem; color: #94A3B8 !important; margin-top: .2rem; }

.ccp-box {
    background: rgba(255,255,255,.05);
    border: 1px solid rgba(255,255,255,.15);
    border-radius: 12px;
    padding: .9rem 1rem;
    font-family: 'Courier New', monospace;
    font-size: .82rem;
    color: #7DD3FC !important;
    margin: .5rem 0;
    word-break: break-all;
}

/* ── Footer ── */
.footer {
    text-align: center;
    color: var(--sub);
    font-size: .82rem;
    padding: 2.5rem 0 1rem;
    opacity: .75;
}

/* ── Animations ── */
@keyframes fadeSlideDown {
    from { opacity:0; transform:translateY(-22px); }
    to   { opacity:1; transform:translateY(0);     }
}
@keyframes fadeSlideUp {
    from { opacity:0; transform:translateY(18px); }
    to   { opacity:1; transform:translateY(0);    }
}

/* ── Divider ── */
.fancy-divider {
    height: 3px;
    background: linear-gradient(90deg, transparent, var(--primary), var(--secondary), transparent);
    border-radius: 3px;
    margin: 1.8rem 0;
    opacity: .4;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────────
if "counter" not in st.session_state:
    st.session_state.counter = 0
if "results" not in st.session_state:
    st.session_state.results = []

# ─────────────────────────────────────────────
#  GEMINI SETUP
# ─────────────────────────────────────────────
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    model = genai.GenerativeModel("gemini-2.0-flash")
except Exception:
    st.error("⚠️  تعذّر الاتصال بـ Gemini API. تأكد من ضبط GEMINI_API_KEY في ملف secrets.toml")
    st.stop()

# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🎓 مولّد الاختبارات")
    st.markdown("---")

    st.markdown("### 📊 الإحصائيات")
    st.markdown(f"""
    <div class="stat-box">
        <div class="stat-number">{st.session_state.counter}</div>
        <div class="stat-label">اختبار مُولَّد منذ بدء الجلسة</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ☕ دعم استمرارية المنصة")
    st.markdown("""
يمكنك دعم هذا المشروع المفتوح عبر **بريدي موب (BaridiMob)**
أو الحساب البريدي الجاري **(CCP)**:
""")
    st.markdown("""
<div class="ccp-box">
  💳 CCP: 0021456789 / clé 32<br>
  🏦 RIP: 00799999002145678932
</div>
""", unsafe_allow_html=True)
    st.caption("شكراً جزيلاً على دعمك 🙏")

    st.markdown("---")
    st.markdown("### 📖 كيفية الاستخدام")
    st.markdown("""
1. اختر المستوى والفصل والمادة
2. أدخل عناوين الدروس
3. اضغط زر التوليد
4. حمّل النتائج بصيغة Word
""")

# ─────────────────────────────────────────────
#  HERO
# ─────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <h1>🏫 مُولِّد الاختبارات الذكي</h1>
    <p>أداة مجانية للأساتذة الجزائريين • الطور الابتدائي • مدعوم بالذكاء الاصطناعي</p>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  DATA
# ─────────────────────────────────────────────
LEVELS = ["السنة الأولى ابتدائي", "السنة الثانية ابتدائي",
          "السنة الثالثة ابتدائي", "السنة الرابعة ابتدائي",
          "السنة الخامسة ابتدائي"]

TRIMESTERS = ["الفصل الأول", "الفصل الثاني", "الفصل الثالث"]

SUBJECTS = [
    "اللغة العربية", "التربية الإسلامية", "التربية المدنية",
    "الرياضيات", "النشاطات العلمية والتكنولوجية",
    "التاريخ والجغرافيا", "التربية التشكيلية",
    "اللغة الفرنسية", "اللغة الأمازيغية",
]

FRENCH_SUBJECTS = ["اللغة الفرنسية"]

# ─────────────────────────────────────────────
#  INPUT SECTION
# ─────────────────────────────────────────────
st.markdown('<div class="fancy-divider"></div>', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""<div class="card"><h3>🎒 المستوى الدراسي</h3>""", unsafe_allow_html=True)
    level = st.selectbox("اختر السنة الدراسية", LEVELS, label_visibility="collapsed")
    st.markdown("</div>", unsafe_allow_html=True)

with col2:
    st.markdown("""<div class="card"><h3>📅 الفصل الدراسي</h3>""", unsafe_allow_html=True)
    trimester = st.selectbox("اختر الفصل", TRIMESTERS, label_visibility="collapsed")
    st.markdown("</div>", unsafe_allow_html=True)

with col3:
    st.markdown("""<div class="card"><h3>📚 المادة المقررة</h3>""", unsafe_allow_html=True)
    subject = st.selectbox("اختر المادة", SUBJECTS, label_visibility="collapsed")
    st.markdown("</div>", unsafe_allow_html=True)

st.markdown("""<div class="card"><h3>📝 عناوين الدروس المُدرَّسة في هذا الفصل</h3>""",
            unsafe_allow_html=True)
lessons = st.text_area(
    "أدخل عناوين الدروس",
    placeholder="مثال:\n• الجمل الاسمية والفعلية\n• علامات الترقيم\n• النعت والمنعوت",
    height=160,
    label_visibility="collapsed",
)
st.markdown("</div>", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  GENERATE BUTTON
# ─────────────────────────────────────────────
generate = st.button("✨ ابدأ التوليد الاحترافي الشامل ✨", type="primary")

# ─────────────────────────────────────────────
#  PROMPT BUILDER
# ─────────────────────────────────────────────
def build_prompt(level, trimester, subject, lessons):
    is_french = subject in FRENCH_SUBJECTS
    lang_note = (
        "يجب أن يكون الاختبار باللغة الفرنسية الأكاديمية الرسمية."
        if is_french else
        "يجب أن تكون جميع نصوص الاختبار مشكولةً شكلاً تاماً وصحيحاً بالحركات لتسهيل قراءة تلميذ الطور الابتدائي."
    )
    header_note = (
        "استخدم الترويسة الرسمية الفرنسية: République Algérienne Démocratique et Populaire / Ministère de l'Éducation Nationale."
        if is_french else
        "استخدم الترويسة الرسمية الجزائرية: الجمهورية الجزائرية الديمقراطية الشعبية / وزارة التربية الوطنية."
    )
    return f"""
أنت مفتش تربوي أول بوزارة التربية الوطنية الجزائرية، متخصص في الطور الابتدائي.

المهمة: أنتج **ثلاثة (3) مقترحات اختبارات منفصلة تماماً** لمادة **{subject}**
للمستوى: **{level}** — **{trimester}**
الدروس المقررة:
{lessons}

المعايير الواجب احترامها:
1. {header_note}
2. يشتمل كل اقتراح على: ترويسة رسمية كاملة (المؤسسة، المستوى، المادة، الفصل، الزمن، التاريخ [اتركه فارغاً]) ثم الأسئلة الموزعة على 10 نقاط، ثم دليل التصحيح النموذجي.
3. {lang_note}
4. تنوّع الأسئلة: فهم، تطبيق، إنتاج، تصحيح أخطاء.
5. تدرّج الصعوبة من السهل إلى الصعب.
6. في أسفل كل اقتراح أضف إطاراً بالنص التالي (لا تغيّر الصياغة):
   ---
   ملاحظات السيد المفتش التربوي: ............................................
   توقيع الأستاذ: .................................  التاريخ: .................
   ---

الصياغة الفاصلة بين الاقتراحات:
افصل بين كل اقتراح بسطر يحتوي فقط على: === الاقتراح الأول === أو === الاقتراح الثاني === إلخ.

ابدأ مباشرةً بأول اقتراح دون مقدمة.
""".strip()

# ─────────────────────────────────────────────
#  DOCX HELPERS
# ─────────────────────────────────────────────
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    pPr_rpr = OxmlElement('w:rPr')
    rtl_el = OxmlElement('w:rtl')
    pPr_rpr.append(rtl_el)
    pPr.append(pPr_rpr)

def build_docx(proposals, level, trimester, subject):
    doc = Document()
    # Page direction
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)

    is_rtl = subject not in FRENCH_SUBJECTS
    align = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT

    for i, text in enumerate(proposals, 1):
        if i > 1:
            doc.add_page_break()
        heading = doc.add_heading(f"الاقتراح رقم {i}", level=1)
        heading.alignment = align
        if is_rtl:
            set_rtl(heading)

        for line in text.strip().splitlines():
            p = doc.add_paragraph(line)
            p.alignment = align
            if is_rtl:
                set_rtl(p)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Simplified Arabic" if is_rtl else "Times New Roman"
            run.font.size = Pt(13)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ─────────────────────────────────────────────
#  GENERATION LOGIC
# ─────────────────────────────────────────────
if generate:
    if not lessons.strip():
        st.warning("⚠️  الرجاء إدخال عناوين الدروس أولاً.")
    else:
        progress_bar = st.progress(0)
        status_text  = st.empty()

        steps = [
            (10,  "🔗  جاري الاتصال بالسيرفر..."),
            (25,  "📡  تم الاتصال! جاري تحليل المنهج الدراسي..."),
            (45,  "🧠  الذكاء الاصطناعي يصوغ الأسئلة بعناية..."),
            (65,  "✍️  جاري مشكلة النصوص وضبط الحركات..."),
            (80,  "📋  جاري بناء دليل التصحيح النموذجي..."),
            (92,  "💅  اللمسات الأخيرة والتنسيق النهائي..."),
        ]

        for pct, msg in steps:
            progress_bar.progress(pct)
            status_text.markdown(f"**{msg} {pct}%**")
            time.sleep(0.45)

        try:
            prompt   = build_prompt(level, trimester, subject, lessons)
            response = model.generate_content(prompt)
            raw      = response.text
        except Exception as e:
            st.error(f"حدث خطأ أثناء الاتصال بـ Gemini: {e}")
            st.stop()

        progress_bar.progress(100)
        status_text.markdown("**✅  اكتمل التوليد بنجاح! 100%**")
        time.sleep(0.5)
        progress_bar.empty()
        status_text.empty()

        # Split proposals
        import re
        parts = re.split(r"===\s*الاقتراح\s+\S+\s*===", raw)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) < 2:
            parts = [raw]  # fallback: show everything as one block

        st.session_state.results  = parts
        st.session_state.counter += len(parts)

# ─────────────────────────────────────────────
#  DISPLAY RESULTS
# ─────────────────────────────────────────────
if st.session_state.results:
    st.markdown('<div class="fancy-divider"></div>', unsafe_allow_html=True)
    st.markdown("## 📄 نتائج التوليد")

    labels = ["الاقتراح الأول 🥇", "الاقتراح الثاني 🥈", "الاقتراح الثالث 🥉"]

    for idx, proposal in enumerate(st.session_state.results):
        label = labels[idx] if idx < len(labels) else f"الاقتراح {idx+1}"
        st.markdown(f"""
        <div class="result-card">
            <span class="result-label">{label}</span><br>
            {proposal.replace(chr(10), '<br>')}
        </div>
        """, unsafe_allow_html=True)

        # Individual download
        docx_single = build_docx([proposal], level, trimester, subject)
        st.download_button(
            label=f"⬇️  تحميل {label} — Word",
            data=docx_single,
            file_name=f"اقتراح_{idx+1}_{subject}_{level}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_{idx}",
        )

    # Bundle download
    st.markdown('<div class="fancy-divider"></div>', unsafe_allow_html=True)
    docx_all = build_docx(st.session_state.results, level, trimester, subject)
    st.download_button(
        label="📦  تحميل الحزمة الكاملة (الاقتراحات الثلاثة) — ملف Word واحد",
        data=docx_all,
        file_name=f"حزمة_اختبارات_{subject}_{level}_{trimester}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_all",
        type="primary",
    )

# ─────────────────────────────────────────────
#  FOOTER
# ─────────────────────────────────────────────
st.markdown("""
<div class="footer">
    <p>صُنع بكل حب لخدمة التعليم في الجزائر 🇩🇿 بواسطة <strong>Laguel Yacine</strong></p>
    <p style="font-size:.75rem; opacity:.6;">مشروع مفتوح المصدر · مرخّص بموجب MIT · 2025</p>
</div>
""", unsafe_allow_html=True)
